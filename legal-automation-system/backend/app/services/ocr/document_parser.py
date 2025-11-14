"""
문서 파싱 및 OCR 처리 시스템
"""

import os
import re
import asyncio
from typing import Dict, Any, List, Optional, Tuple
from pathlib import Path
import hashlib
from datetime import datetime

import cv2
import numpy as np
from PIL import Image
import pytesseract
import easyocr
import pdfplumber
from pdf2image import convert_from_path, convert_from_bytes
import docx
import hwpapi

from app.core.config import settings
from app.models.document import Document, DocumentType
from app.core.logging import logger


class OCRProcessor:
    """OCR 처리 엔진"""

    def __init__(self):
        """OCR 엔진 초기화"""
        # EasyOCR 초기화 (한국어 + 영어)
        self.easyocr_reader = easyocr.Reader(['ko', 'en'], gpu=settings.OCR_USE_GPU)

        # Tesseract 설정
        if settings.TESSERACT_PATH:
            pytesseract.pytesseract.tesseract_cmd = settings.TESSERACT_PATH

        # OCR 설정
        self.confidence_threshold = 0.5
        self.language = 'kor+eng'  # Tesseract 한국어+영어

    async def process_image(self, image: np.ndarray, engine: str = "auto") -> Dict[str, Any]:
        """
        이미지에서 텍스트 추출

        Args:
            image: 입력 이미지 (numpy array)
            engine: OCR 엔진 (auto, tesseract, easyocr)

        Returns:
            추출된 텍스트 및 메타데이터
        """
        try:
            # 이미지 전처리
            processed_image = await self._preprocess_image(image)

            if engine == "auto":
                # 자동 엔진 선택 (이미지 특성에 따라)
                engine = self._select_best_engine(processed_image)

            # OCR 수행
            if engine == "tesseract":
                result = await self._ocr_tesseract(processed_image)
            elif engine == "easyocr":
                result = await self._ocr_easyocr(processed_image)
            else:
                # 두 엔진 모두 사용하여 최적 결과 선택
                result = await self._ocr_hybrid(processed_image)

            # 후처리
            result['text'] = await self._postprocess_text(result['text'])

            return result

        except Exception as e:
            logger.error(f"OCR processing error: {e}")
            return {
                'text': '',
                'confidence': 0.0,
                'engine': engine,
                'error': str(e)
            }

    async def _preprocess_image(self, image: np.ndarray) -> np.ndarray:
        """이미지 전처리"""
        # 그레이스케일 변환
        if len(image.shape) == 3:
            gray = cv2.cvtColor(image, cv2.COLOR_BGR2GRAY)
        else:
            gray = image

        # 노이즈 제거
        denoised = cv2.fastNlMeansDenoising(gray)

        # 이진화
        _, binary = cv2.threshold(denoised, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)

        # 스큐 보정
        corrected = await self._correct_skew(binary)

        return corrected

    async def _correct_skew(self, image: np.ndarray) -> np.ndarray:
        """이미지 기울기 보정"""
        # 엣지 검출
        edges = cv2.Canny(image, 50, 150, apertureSize=3)

        # 허프 변환으로 직선 검출
        lines = cv2.HoughLines(edges, 1, np.pi/180, 200)

        if lines is not None:
            # 각도 계산
            angles = []
            for rho, theta in lines[:, 0]:
                angle = np.degrees(theta) - 90
                if -45 <= angle <= 45:
                    angles.append(angle)

            if angles:
                median_angle = np.median(angles)

                # 회전 행렬
                (h, w) = image.shape[:2]
                center = (w // 2, h // 2)
                M = cv2.getRotationMatrix2D(center, median_angle, 1.0)

                # 이미지 회전
                rotated = cv2.warpAffine(image, M, (w, h),
                                        flags=cv2.INTER_CUBIC,
                                        borderMode=cv2.BORDER_REPLICATE)
                return rotated

        return image

    def _select_best_engine(self, image: np.ndarray) -> str:
        """이미지 특성에 따라 최적 OCR 엔진 선택"""
        # 이미지 특성 분석
        height, width = image.shape[:2]

        # 고해상도 이미지는 EasyOCR이 더 정확
        if width > 2000 or height > 2000:
            return "easyocr"

        # 단순한 텍스트는 Tesseract가 빠름
        return "tesseract"

    async def _ocr_tesseract(self, image: np.ndarray) -> Dict[str, Any]:
        """Tesseract OCR 수행"""
        try:
            # OCR 설정
            custom_config = r'--oem 3 --psm 6'

            # 텍스트 추출
            text = pytesseract.image_to_string(
                image,
                lang=self.language,
                config=custom_config
            )

            # 상세 정보 추출
            data = pytesseract.image_to_data(
                image,
                lang=self.language,
                output_type=pytesseract.Output.DICT
            )

            # 신뢰도 계산
            confidences = [float(c) for c in data['conf'] if c != '-1']
            avg_confidence = sum(confidences) / len(confidences) if confidences else 0

            return {
                'text': text,
                'confidence': avg_confidence / 100,
                'engine': 'tesseract',
                'details': data
            }

        except Exception as e:
            logger.error(f"Tesseract OCR error: {e}")
            raise

    async def _ocr_easyocr(self, image: np.ndarray) -> Dict[str, Any]:
        """EasyOCR 수행"""
        try:
            # 텍스트 추출
            results = self.easyocr_reader.readtext(image)

            # 결과 파싱
            text_lines = []
            confidences = []

            for (bbox, text, prob) in results:
                text_lines.append(text)
                confidences.append(prob)

            full_text = '\n'.join(text_lines)
            avg_confidence = sum(confidences) / len(confidences) if confidences else 0

            return {
                'text': full_text,
                'confidence': avg_confidence,
                'engine': 'easyocr',
                'details': results
            }

        except Exception as e:
            logger.error(f"EasyOCR error: {e}")
            raise

    async def _ocr_hybrid(self, image: np.ndarray) -> Dict[str, Any]:
        """하이브리드 OCR (두 엔진 결합)"""
        # 병렬로 두 엔진 실행
        tasks = [
            self._ocr_tesseract(image),
            self._ocr_easyocr(image)
        ]

        results = await asyncio.gather(*tasks, return_exceptions=True)

        # 오류가 없는 결과 중 신뢰도가 높은 것 선택
        valid_results = [r for r in results if not isinstance(r, Exception)]

        if not valid_results:
            raise Exception("All OCR engines failed")

        # 신뢰도가 가장 높은 결과 선택
        best_result = max(valid_results, key=lambda x: x['confidence'])
        best_result['engine'] = 'hybrid'

        return best_result

    async def _postprocess_text(self, text: str) -> str:
        """추출된 텍스트 후처리"""
        if not text:
            return text

        # 불필요한 공백 제거
        text = re.sub(r'\s+', ' ', text)

        # 줄바꿈 정리
        text = re.sub(r'\n{3,}', '\n\n', text)

        # 특수문자 정리
        text = text.replace('｜', '|')
        text = text.replace('－', '-')

        # 한글 조사 수정 (OCR 오류 보정)
        corrections = {
            '을를': '을',
            '는은': '는',
            '이가': '이',
            '와과': '와'
        }

        for wrong, correct in corrections.items():
            text = text.replace(wrong, correct)

        return text.strip()


class DocumentParser:
    """문서 파싱 및 구조 분석"""

    def __init__(self):
        """문서 파서 초기화"""
        self.ocr_processor = OCRProcessor()
        self.supported_formats = {
            '.pdf': self._parse_pdf,
            '.jpg': self._parse_image,
            '.jpeg': self._parse_image,
            '.png': self._parse_image,
            '.tiff': self._parse_image,
            '.docx': self._parse_docx,
            '.hwp': self._parse_hwp,
            '.txt': self._parse_text
        }

    async def parse_document(
        self,
        file_path: str,
        extract_metadata: bool = True,
        extract_structure: bool = True
    ) -> Dict[str, Any]:
        """
        문서 파싱 및 텍스트 추출

        Args:
            file_path: 문서 파일 경로
            extract_metadata: 메타데이터 추출 여부
            extract_structure: 구조 정보 추출 여부

        Returns:
            파싱된 문서 정보
        """
        try:
            file_path = Path(file_path)

            if not file_path.exists():
                raise FileNotFoundError(f"File not found: {file_path}")

            # 파일 확장자 확인
            ext = file_path.suffix.lower()
            if ext not in self.supported_formats:
                raise ValueError(f"Unsupported file format: {ext}")

            # 문서 파싱
            parse_func = self.supported_formats[ext]
            result = await parse_func(file_path)

            # 메타데이터 추출
            if extract_metadata:
                result['metadata'] = await self._extract_metadata(file_path, result)

            # 구조 분석
            if extract_structure:
                result['structure'] = await self._analyze_structure(result['text'])

            # 문서 타입 분류
            result['document_type'] = await self._classify_document(result['text'])

            # 해시 생성 (중복 검사용)
            result['hash'] = self._generate_hash(result['text'])

            return result

        except Exception as e:
            logger.error(f"Document parsing error: {e}")
            raise

    async def _parse_pdf(self, file_path: Path) -> Dict[str, Any]:
        """PDF 문서 파싱"""
        text_pages = []
        ocr_pages = []
        total_pages = 0

        # pdfplumber로 텍스트 추출 시도
        with pdfplumber.open(file_path) as pdf:
            total_pages = len(pdf.pages)

            for i, page in enumerate(pdf.pages):
                # 텍스트 추출
                page_text = page.extract_text() or ''

                if page_text.strip():
                    # 텍스트가 있으면 저장
                    text_pages.append({
                        'page': i + 1,
                        'text': page_text,
                        'method': 'extraction'
                    })
                else:
                    # 텍스트가 없으면 OCR 대상
                    ocr_pages.append(i)

        # OCR 필요한 페이지 처리
        if ocr_pages:
            # PDF를 이미지로 변환
            images = convert_from_path(str(file_path), dpi=300)

            for page_num in ocr_pages:
                if page_num < len(images):
                    # 이미지를 numpy 배열로 변환
                    image = np.array(images[page_num])

                    # OCR 수행
                    ocr_result = await self.ocr_processor.process_image(image)

                    text_pages.append({
                        'page': page_num + 1,
                        'text': ocr_result['text'],
                        'method': 'ocr',
                        'confidence': ocr_result['confidence']
                    })

        # 페이지 순서대로 정렬
        text_pages.sort(key=lambda x: x['page'])

        # 전체 텍스트 결합
        full_text = '\n\n'.join([p['text'] for p in text_pages])

        return {
            'text': full_text,
            'pages': text_pages,
            'total_pages': total_pages,
            'format': 'pdf'
        }

    async def _parse_image(self, file_path: Path) -> Dict[str, Any]:
        """이미지 파일 파싱"""
        # 이미지 로드
        image = cv2.imread(str(file_path))

        if image is None:
            raise ValueError(f"Failed to load image: {file_path}")

        # OCR 수행
        ocr_result = await self.ocr_processor.process_image(image)

        return {
            'text': ocr_result['text'],
            'confidence': ocr_result['confidence'],
            'engine': ocr_result['engine'],
            'format': 'image'
        }

    async def _parse_docx(self, file_path: Path) -> Dict[str, Any]:
        """DOCX 문서 파싱"""
        doc = docx.Document(file_path)

        # 텍스트 추출
        paragraphs = []
        for para in doc.paragraphs:
            if para.text.strip():
                paragraphs.append(para.text)

        # 표 텍스트 추출
        tables = []
        for table in doc.tables:
            table_text = []
            for row in table.rows:
                row_text = [cell.text for cell in row.cells]
                table_text.append(' | '.join(row_text))
            tables.append('\n'.join(table_text))

        # 전체 텍스트 결합
        full_text = '\n\n'.join(paragraphs)
        if tables:
            full_text += '\n\n' + '\n\n'.join(tables)

        return {
            'text': full_text,
            'format': 'docx',
            'paragraphs': len(paragraphs),
            'tables': len(tables)
        }

    async def _parse_hwp(self, file_path: Path) -> Dict[str, Any]:
        """HWP 문서 파싱 (한글 파일)"""
        try:
            # hwpapi를 사용하여 텍스트 추출
            hwp = hwpapi.open(str(file_path))
            text = hwp.get_text()
            hwp.close()

            return {
                'text': text,
                'format': 'hwp'
            }
        except Exception as e:
            logger.warning(f"HWP parsing failed, trying OCR: {e}")

            # HWP 파싱 실패 시 PDF로 변환 후 시도
            # (실제 구현에서는 HWP to PDF 변환 도구 필요)
            raise NotImplementedError("HWP OCR fallback not implemented")

    async def _parse_text(self, file_path: Path) -> Dict[str, Any]:
        """텍스트 파일 파싱"""
        # 인코딩 감지 및 읽기
        encodings = ['utf-8', 'euc-kr', 'cp949']

        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as f:
                    text = f.read()
                    return {
                        'text': text,
                        'format': 'text',
                        'encoding': encoding
                    }
            except UnicodeDecodeError:
                continue

        raise ValueError(f"Unable to decode text file: {file_path}")

    async def _extract_metadata(self, file_path: Path, parse_result: Dict[str, Any]) -> Dict[str, Any]:
        """문서 메타데이터 추출"""
        metadata = {
            'filename': file_path.name,
            'size': file_path.stat().st_size,
            'created': datetime.fromtimestamp(file_path.stat().st_ctime),
            'modified': datetime.fromtimestamp(file_path.stat().st_mtime),
            'format': parse_result.get('format'),
            'pages': parse_result.get('total_pages', 1)
        }

        # 텍스트 통계
        text = parse_result.get('text', '')
        metadata['statistics'] = {
            'characters': len(text),
            'words': len(text.split()),
            'lines': len(text.splitlines())
        }

        return metadata

    async def _analyze_structure(self, text: str) -> Dict[str, Any]:
        """문서 구조 분석"""
        structure = {
            'sections': [],
            'clauses': [],
            'signatures': [],
            'dates': []
        }

        # 섹션 찾기 (제1조, 제2조 등)
        clause_pattern = r'제(\d+)조\s*\(([^)]+)\)'
        clauses = re.finditer(clause_pattern, text)

        for match in clauses:
            structure['clauses'].append({
                'number': match.group(1),
                'title': match.group(2),
                'position': match.start()
            })

        # 서명란 찾기
        signature_patterns = [
            r'(갑|을)\s*:\s*([^\n]+)',
            r'(매도인|매수인)\s*:\s*([^\n]+)',
            r'(임대인|임차인)\s*:\s*([^\n]+)'
        ]

        for pattern in signature_patterns:
            signatures = re.finditer(pattern, text)
            for match in signatures:
                structure['signatures'].append({
                    'role': match.group(1),
                    'name': match.group(2).strip(),
                    'position': match.start()
                })

        # 날짜 찾기
        date_pattern = r'\d{4}년\s*\d{1,2}월\s*\d{1,2}일'
        dates = re.finditer(date_pattern, text)

        for match in dates:
            structure['dates'].append({
                'date': match.group(),
                'position': match.start()
            })

        return structure

    async def _classify_document(self, text: str) -> str:
        """문서 타입 분류"""
        # 키워드 기반 분류
        classifications = {
            DocumentType.CONTRACT: ['계약서', '계약', '약정', '갑', '을'],
            DocumentType.LAWSUIT: ['소장', '원고', '피고', '청구', '소송'],
            DocumentType.NOTICE: ['내용증명', '통지', '최고장'],
            DocumentType.AGREEMENT: ['합의서', '협의', '합의'],
            DocumentType.APPLICATION: ['신청서', '신청인', '신청'],
            DocumentType.PETITION: ['탄원서', '탄원', '청원'],
            DocumentType.CERTIFICATE: ['증명서', '확인서', '증명']
        }

        scores = {}
        for doc_type, keywords in classifications.items():
            score = sum(1 for keyword in keywords if keyword in text)
            if score > 0:
                scores[doc_type] = score

        if scores:
            # 가장 높은 점수의 타입 반환
            return max(scores, key=scores.get).value

        return DocumentType.OTHER.value

    def _generate_hash(self, text: str) -> str:
        """텍스트 해시 생성"""
        return hashlib.sha256(text.encode('utf-8')).hexdigest()


class DocumentProcessor:
    """문서 처리 통합 시스템"""

    def __init__(self):
        """문서 프로세서 초기화"""
        self.parser = DocumentParser()
        self.storage_path = Path(settings.DOCUMENT_STORAGE_PATH)
        self.storage_path.mkdir(parents=True, exist_ok=True)

    async def process_document(
        self,
        file_path: str,
        user_id: int,
        save_to_db: bool = True
    ) -> Document:
        """
        문서 처리 및 저장

        Args:
            file_path: 처리할 문서 경로
            user_id: 사용자 ID
            save_to_db: DB 저장 여부

        Returns:
            처리된 문서 객체
        """
        try:
            # 문서 파싱
            parsed = await self.parser.parse_document(file_path)

            # 문서 저장
            if save_to_db:
                document = await self._save_document(parsed, user_id)

                # 원본 파일 저장
                await self._store_original(file_path, document.id)

                # 처리된 텍스트 저장
                await self._store_processed(parsed['text'], document.id)

                return document

            return parsed

        except Exception as e:
            logger.error(f"Document processing error: {e}")
            raise

    async def _save_document(self, parsed: Dict[str, Any], user_id: int) -> Document:
        """문서 정보 DB 저장"""
        # TODO: 실제 DB 저장 로직 구현
        pass

    async def _store_original(self, file_path: str, document_id: int):
        """원본 파일 저장"""
        source = Path(file_path)
        dest = self.storage_path / 'originals' / f"{document_id}{source.suffix}"
        dest.parent.mkdir(parents=True, exist_ok=True)

        # 파일 복사
        import shutil
        shutil.copy2(source, dest)

    async def _store_processed(self, text: str, document_id: int):
        """처리된 텍스트 저장"""
        dest = self.storage_path / 'processed' / f"{document_id}.txt"
        dest.parent.mkdir(parents=True, exist_ok=True)

        with open(dest, 'w', encoding='utf-8') as f:
            f.write(text)