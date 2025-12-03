"""
Draft Service - Business logic for draft generation with RAG
Orchestrates Qdrant RAG + OpenAI GPT-4o for draft preview
"""

from sqlalchemy.orm import Session
from typing import List, Tuple
from datetime import datetime, timezone
from io import BytesIO

from app.db.schemas import (
    DraftPreviewRequest,
    DraftPreviewResponse,
    DraftCitation,
    DraftExportFormat
)
from app.repositories.case_repository import CaseRepository
from app.repositories.case_member_repository import CaseMemberRepository
from app.utils.dynamo import get_evidence_by_case
from app.utils.qdrant import search_evidence_by_semantic, search_legal_knowledge
from app.utils.openai_client import generate_chat_completion
from app.middleware import NotFoundError, PermissionError, ValidationError

# Optional: python-docx for DOCX generation
try:
    from docx import Document
    from docx.shared import Pt, Inches  # noqa: F401
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


class DraftService:
    """
    Service for draft generation with RAG
    """

    def __init__(self, db: Session):
        self.db = db
        self.case_repo = CaseRepository(db)
        self.member_repo = CaseMemberRepository(db)

    def generate_draft_preview(
        self,
        case_id: str,
        request: DraftPreviewRequest,
        user_id: str
    ) -> DraftPreviewResponse:
        """
        Generate draft preview using RAG + GPT-4o

        Process:
        1. Validate case access
        2. Retrieve evidence metadata from DynamoDB
        3. Perform semantic search in Qdrant (RAG)
        4. Build GPT-4o prompt with RAG context
        5. Generate draft text
        6. Extract citations

        Args:
            case_id: Case ID
            request: Draft generation request (sections, language, style)
            user_id: User ID requesting draft

        Returns:
            Draft preview with citations

        Raises:
            NotFoundError: Case not found
            PermissionError: User does not have access to case
            ValidationError: No evidence in case
        """
        # 1. Validate case access
        case = self.case_repo.get_by_id(case_id)
        if not case:
            raise NotFoundError("Case")

        if not self.member_repo.has_access(case_id, user_id):
            raise PermissionError("You do not have access to this case")

        # 2. Retrieve evidence metadata from DynamoDB
        evidence_list = get_evidence_by_case(case_id)

        # Check if there's any evidence
        if not evidence_list:
            raise ValidationError("사건에 증거가 하나도 없습니다. 증거를 업로드한 후 초안을 생성해 주세요.")

        # Filter for completed evidence only (status="done")
        # Note: Currently filtering for reference, may be used for future enhancements
        _ = [ev for ev in evidence_list if ev.get("status") == "done"]

        # 3. Perform semantic RAG search in Qdrant (evidence + legal)
        rag_results = self._perform_rag_search(case_id, request.sections)
        evidence_results = rag_results.get("evidence", [])
        legal_results = rag_results.get("legal", [])

        # 4. Build GPT-4o prompt with RAG context
        prompt_messages = self._build_draft_prompt(
            case=case,
            sections=request.sections,
            evidence_context=evidence_results,
            legal_context=legal_results,
            language=request.language,
            style=request.style
        )

        # 5. Generate draft text using GPT-4o
        draft_text = generate_chat_completion(
            messages=prompt_messages,
            temperature=0.3,  # Low temperature for consistent legal writing
            max_tokens=4000
        )

        # 6. Extract citations from RAG results
        citations = self._extract_citations(evidence_results)

        return DraftPreviewResponse(
            case_id=case_id,
            draft_text=draft_text,
            citations=citations,
            generated_at=datetime.now(timezone.utc)
        )

    def _perform_rag_search(self, case_id: str, sections: List[str]) -> dict:
        """
        Perform semantic search in Qdrant for RAG context

        Args:
            case_id: Case ID
            sections: Sections being generated

        Returns:
            Dict with 'evidence' and 'legal' results
        """
        # Build search query based on sections
        if "청구원인" in sections:
            # Search for fault evidence (guilt factors)
            query = "이혼 사유 귀책사유 폭언 불화 부정행위"
            evidence_results = search_evidence_by_semantic(
                case_id=case_id,
                query=query,
                top_k=10
            )
            # Search legal knowledge for divorce grounds
            legal_results = search_legal_knowledge(
                query="재판상 이혼 사유 민법 제840조",
                top_k=5,
                doc_type="statute"
            )
        else:
            # General search for all sections
            query = " ".join(sections)
            evidence_results = search_evidence_by_semantic(
                case_id=case_id,
                query=query,
                top_k=5
            )
            legal_results = search_legal_knowledge(
                query="이혼 " + query,
                top_k=3
            )

        return {
            "evidence": evidence_results,
            "legal": legal_results
        }

    def _build_draft_prompt(
        self,
        case: any,
        sections: List[str],
        evidence_context: List[dict],
        legal_context: List[dict],
        language: str,
        style: str
    ) -> List[dict]:
        """
        Build GPT-4o prompt with evidence and legal RAG context

        Args:
            case: Case object
            sections: Sections to generate
            evidence_context: Evidence RAG search results
            legal_context: Legal knowledge RAG search results
            language: Language (ko/en)
            style: Writing style

        Returns:
            List of messages for GPT-4o
        """
        # System message - define role and constraints
        system_message = {
            "role": "system",
            "content": """당신은 대한민국 가정법원에 제출하는 정식 소장(訴狀)을 작성하는 전문 법률가입니다.

## 필수 출력 형식
- **마크다운 형식**으로 출력하세요
- 각 섹션은 반드시 빈 줄로 구분
- 제목은 ## 또는 ### 사용
- 표는 마크다운 테이블 형식 사용

---

# 소    장

**사건명:** 이혼 등 청구의 소

---

## 【원  고】
- 성명: [원고 성명]
- 주민등록번호: ○○○○○○-○******
- 주소: [상세 주소]
- 등록기준지: [등록기준지]

## 【피  고】
- 성명: [피고 성명]
- 주민등록번호: ○○○○○○-○******
- 주소: [상세 주소]
- 등록기준지: [등록기준지]

---

# 청 구 취 지

1. 원고와 피고는 이혼한다.

2. 피고는 원고에게 위자료로 금 ○○,○○○,○○○원 및 이에 대하여 **이 사건 소장 부본 송달일 다음날부터 다 갚는 날까지 연 12%의 비율**로 계산한 지연손해금을 지급하라.

3. 피고는 원고에게 재산분할로 금 ○○○,○○○,○○○원을 지급하라.

4. 사건본인 [자녀명](○○○○년 ○월 ○일생)의 친권자 및 양육자로 원고를 지정한다.

5. 피고는 원고에게 사건본인의 양육비로 이 판결 확정일로부터 사건본인이 성년에 이르기까지 매월 말일에 금 ○○○,○○○원씩 지급하라.
   > ※ 양육비산정기준표 기준

6. 소송비용은 피고의 부담으로 한다.

7. 제2, 3, 5항은 가집행할 수 있다.

**라는 판결을 구합니다.**

---

# 청 구 원 인

## 제1. 당사자들의 관계

원고와 피고는 **○○○○년 ○월 ○일** 혼인신고를 마친 법률상 부부로서, 슬하에 사건본인 [자녀명](○○○○년 ○월 ○일생)을 두고 있습니다.

## 제2. 혼인생활의 경과

### 가. 혼인 초기 (○○○○년 ~ ○○○○년)
[구체적인 혼인생활 내용]

### 나. 갈등의 시작 (○○○○년 ○월경)
[갈등 발생 시점과 원인]

## 제3. 이혼 사유 (민법 제840조)

### 가. 법적 근거
피고의 아래 행위로 혼인관계가 회복할 수 없을 정도로 파탄되었으므로, **민법 제840조 제6호 "혼인을 계속하기 어려운 중대한 사유"**에 해당합니다.

### 나. 피고의 유책행위

| 일시 | 내용 | 증거 |
|------|------|------|
| ○○○○.○○.○○ | [구체적 사실] | [갑 제○호증] |
| ○○○○.○○.○○ | [구체적 사실] | [갑 제○호증] |

> 📌 **증거에서 확인된 내용:**
> - "[증거에서 발췌한 구체적 발언/행위]"

## 제4. 위자료 청구

### 가. 청구금액: 금 **○○,○○○,○○○원**

### 나. 산정근거
| 항목 | 내용 |
|------|------|
| 혼인기간 | ○○년 ○개월 |
| 유책행위 정도 | [구체적 기술] |
| 정신적 고통 | [구체적 기술] |

### 다. 지연손해금
소송촉진 등에 관한 특례법 제3조에 따라 소장 부본 송달일 다음날부터 **연 12%** 지연손해금

## 제5. 재산분할 청구

### 가. 분할대상 재산
| 재산 종류 | 명의자 | 현재 가액 |
|-----------|--------|-----------|
| [부동산/예금] | [원고/피고] | ○○○,○○○원 |

### 나. 기여도: 원고 **○○%** / 피고 **○○%**

## 제6. 친권자·양육자 지정 및 양육비

### 가. 친권자·양육자로 원고 지정 사유
- 사건본인 연령: ○○세
- 양육환경: [주거, 경제력]

### 나. 양육비 산정 (양육비산정기준표 기준)
| 항목 | 내용 |
|------|------|
| 부모 합산소득 | 월 ○○○만원 |
| 자녀 연령 | ○○세 |
| **청구금액** | **월 ○○○,○○○원** |

---

# 입 증 방 법

| 증거번호 | 증거명 | 요지 |
|----------|--------|------|
| 갑 제1호증 | [증거명] | [날짜, 내용] |
| 갑 제2호증 | [증거명] | [날짜, 내용] |
| 갑 제3호증 | 혼인관계증명서 | - |
| 갑 제4호증 | 가족관계증명서 | - |

---

# 첨 부 서 류

1. 위 입증방법 각 1통
2. 소장부본 1통
3. 송달료납부서 1통

---

**○○○○년  ○○월  ○○일**

**위 원고  [원고 성명]  (인)**

**○○가정법원 귀중**

---

## 작성 원칙

1. **증거 인용**: `[갑 제N호증]` 형식으로 정확히 인용
2. **법률 근거**: 민법 제840조 명시적 인용
3. **구체성**: 날짜, 발언 내용을 증거에서 직접 인용
4. **금액 산정 근거**:
   - 위자료: 혼인기간, 유책정도
   - 양육비: 양육비산정기준표 명시
   - 지연손해금: 연 12% (소송촉진법)
5. **Placeholder**: 확인 안 된 정보는 `○○○` 형식

⚠️ **본 문서는 AI가 생성한 초안이며, 변호사의 검토 및 수정이 필수입니다.**
"""
        }

        # Build context strings
        evidence_context_str = self._format_evidence_context(evidence_context)
        legal_context_str = self._format_legal_context(legal_context)

        # User message - include case info, evidence, and legal context
        user_message = {
            "role": "user",
            "content": f"""
다음 정보를 바탕으로 이혼 소송 소장 초안을 작성해 주세요.

**사건 정보:**
- 사건명: {case.title}
- 사건 설명: {case.description or "N/A"}

**생성할 섹션:**
{", ".join(sections)}

**관련 법률 조문:**
{legal_context_str}

**증거 자료:**
{evidence_context_str}

**요청사항:**
- 언어: {language}
- 스타일: {style}
- 위 법률 조문과 증거를 기반으로 법률적 논리를 구성해 주세요
- 이혼 사유는 반드시 민법 제840조를 인용하여 작성하세요
- 각 주장에 대해 증거 번호를 명시해 주세요 (예: [갑 제1호증], [갑 제2호증])

소장 초안을 작성해 주세요.
"""
        }

        return [system_message, user_message]

    def _format_legal_context(self, legal_results: List[dict]) -> str:
        """
        Format legal knowledge search results for GPT-4o prompt

        Args:
            legal_results: List of legal documents from RAG search

        Returns:
            Formatted legal context string
        """
        if not legal_results:
            return "(관련 법률 조문 없음)"

        context_parts = []
        for doc in legal_results:
            article_number = doc.get("article_number", "")
            statute_name = doc.get("statute_name", "민법")
            # Qdrant payload uses "document" field, not "text"
            content = doc.get("document", "") or doc.get("text", "")

            if article_number and content:
                context_parts.append(f"""
【{statute_name} {article_number}】
{content}
""")

        return "\n".join(context_parts) if context_parts else "(관련 법률 조문 없음)"

    def _format_evidence_context(self, evidence_results: List[dict]) -> str:
        """
        Format evidence search results for GPT-4o prompt

        Args:
            evidence_results: List of evidence documents from RAG search

        Returns:
            Formatted evidence context string
        """
        if not evidence_results:
            return "(증거 자료 없음 - 기본 템플릿으로 작성)"

        context_parts = []
        for i, doc in enumerate(evidence_results, start=1):
            evidence_id = doc.get("id", f"evidence_{i}")
            content = doc.get("content", "")
            labels = doc.get("labels", [])
            speaker = doc.get("speaker", "")
            timestamp = doc.get("timestamp", "")

            # Truncate content if too long
            if len(content) > 500:
                content = content[:500] + "..."

            context_parts.append(f"""
[갑 제{i}호증] (ID: {evidence_id})
- 분류: {", ".join(labels) if labels else "N/A"}
- 화자: {speaker or "N/A"}
- 시점: {timestamp or "N/A"}
- 내용: {content}
""")

        return "\n".join(context_parts)

    def _format_rag_context(self, rag_results: List[dict]) -> str:
        """
        Format RAG search results for GPT-4o prompt

        Args:
            rag_results: List of evidence documents from RAG search

        Returns:
            Formatted context string
        """
        if not rag_results:
            return "(증거 자료 없음 - 기본 템플릿으로 작성)"

        context_parts = []
        for i, doc in enumerate(rag_results, start=1):
            evidence_id = doc.get("id", f"evidence_{i}")
            content = doc.get("content", "")
            labels = doc.get("labels", [])
            speaker = doc.get("speaker", "")
            timestamp = doc.get("timestamp", "")

            # Truncate content if too long
            if len(content) > 500:
                content = content[:500] + "..."

            context_parts.append(f"""
[증거 {i}] (ID: {evidence_id})
- 분류: {", ".join(labels) if labels else "N/A"}
- 화자: {speaker or "N/A"}
- 시점: {timestamp or "N/A"}
- 내용: {content}
""")

        return "\n".join(context_parts)

    def _extract_citations(self, rag_results: List[dict]) -> List[DraftCitation]:
        """
        Extract citations from RAG results

        Args:
            rag_results: List of evidence documents from RAG search

        Returns:
            List of DraftCitation objects
        """
        citations = []

        for doc in rag_results:
            evidence_id = doc.get("evidence_id") or doc.get("id")
            content = doc.get("content", "")
            labels = doc.get("labels", [])

            # Create snippet (first 200 chars)
            snippet = content[:200] + "..." if len(content) > 200 else content

            citations.append(
                DraftCitation(
                    evidence_id=evidence_id,
                    snippet=snippet,
                    labels=labels
                )
            )

        return citations

    def export_draft(
        self,
        case_id: str,
        user_id: str,
        export_format: DraftExportFormat = DraftExportFormat.DOCX
    ) -> Tuple[BytesIO, str, str]:
        """
        Export draft as DOCX or PDF file

        Process:
        1. Validate case access
        2. Generate draft preview using RAG + GPT-4o
        3. Convert to requested format (DOCX or PDF)

        Args:
            case_id: Case ID
            user_id: User ID requesting export
            export_format: Output format (docx or pdf)

        Returns:
            Tuple of (file_bytes, filename, content_type)

        Raises:
            NotFoundError: Case not found
            PermissionError: User does not have access to case
            ValidationError: Export format not supported or missing dependencies
        """
        # 1. Validate case access
        case = self.case_repo.get_by_id(case_id)
        if not case:
            raise NotFoundError("Case")

        if not self.member_repo.has_access(case_id, user_id):
            raise PermissionError("You do not have access to this case")

        # 2. Generate draft preview
        request = DraftPreviewRequest()  # Use default sections
        draft_response = self.generate_draft_preview(case_id, request, user_id)

        # 3. Convert to requested format
        if export_format == DraftExportFormat.DOCX:
            return self._generate_docx(case, draft_response)
        elif export_format == DraftExportFormat.PDF:
            return self._generate_pdf(case, draft_response)
        else:
            raise ValidationError(f"Unsupported export format: {export_format}")

    def _generate_docx(
        self,
        case,
        draft_response: DraftPreviewResponse
    ) -> Tuple[BytesIO, str, str]:
        """
        Generate DOCX file from draft response

        Args:
            case: Case object
            draft_response: Generated draft preview

        Returns:
            Tuple of (file_bytes, filename, content_type)
        """
        if not DOCX_AVAILABLE:
            raise ValidationError(
                "DOCX export is not available. "
                "Please install python-docx: pip install python-docx"
            )

        # Create document
        doc = Document()

        # Title
        title = doc.add_heading("이혼 소송 준비서면 (초안)", level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Case info
        doc.add_paragraph()
        case_info = doc.add_paragraph()
        case_info.add_run(f"사건명: {case.title}").bold = True
        doc.add_paragraph(f"생성일시: {draft_response.generated_at.strftime('%Y-%m-%d %H:%M:%S')}")

        # Draft content
        doc.add_heading("본문", level=1)
        for paragraph_text in draft_response.draft_text.split("\n\n"):
            if paragraph_text.strip():
                doc.add_paragraph(paragraph_text.strip())

        # Citations
        if draft_response.citations:
            doc.add_heading("인용 증거", level=1)
            for i, citation in enumerate(draft_response.citations, 1):
                p = doc.add_paragraph()
                p.add_run(f"[증거 {i}] ").bold = True
                p.add_run(f"(ID: {citation.evidence_id})")
                doc.add_paragraph(f"  - 분류: {', '.join(citation.labels) if citation.labels else 'N/A'}")
                doc.add_paragraph(f"  - 내용: {citation.snippet}")

        # Disclaimer
        doc.add_paragraph()
        disclaimer = doc.add_paragraph()
        disclaimer.add_run(
            "⚠️ 본 문서는 AI가 생성한 초안이며, "
            "변호사의 검토 및 수정이 필수입니다."
        ).italic = True

        # Save to BytesIO
        file_buffer = BytesIO()
        doc.save(file_buffer)
        file_buffer.seek(0)

        # Generate filename
        safe_title = case.title.replace(" ", "_")[:30]
        filename = f"draft_{safe_title}_{draft_response.generated_at.strftime('%Y%m%d')}.docx"
        content_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

        return file_buffer, filename, content_type

    def _generate_pdf(
        self,
        case,
        draft_response: DraftPreviewResponse
    ) -> Tuple[BytesIO, str, str]:
        """
        Generate PDF file from draft response with Korean font support

        Features:
        - A4 layout for legal documents
        - Korean font support (Noto Sans KR or system fallback)
        - Legal document template structure
        - Citations section

        Args:
            case: Case object
            draft_response: Generated draft preview

        Returns:
            Tuple of (file_bytes, filename, content_type)
        """
        try:
            from reportlab.lib.pagesizes import A4
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.platypus import (
                SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
            )
            from reportlab.lib.units import inch, mm
            from reportlab.lib import colors
            from reportlab.pdfbase import pdfmetrics
            from reportlab.pdfbase.ttfonts import TTFont
            from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY
            import os

            # Register Korean font
            korean_font_registered = self._register_korean_font(pdfmetrics, TTFont)
            font_name = "NotoSansKR" if korean_font_registered else "Helvetica"

            file_buffer = BytesIO()

            # A4 with proper margins for legal documents
            doc = SimpleDocTemplate(
                file_buffer,
                pagesize=A4,
                leftMargin=25 * mm,
                rightMargin=25 * mm,
                topMargin=25 * mm,
                bottomMargin=25 * mm
            )

            # Create custom styles with Korean font
            styles = getSampleStyleSheet()

            # Title style
            title_style = ParagraphStyle(
                'KoreanTitle',
                parent=styles['Title'],
                fontName=font_name,
                fontSize=18,
                alignment=TA_CENTER,
                spaceAfter=20
            )

            # Heading style
            heading_style = ParagraphStyle(
                'KoreanHeading',
                parent=styles['Heading2'],
                fontName=font_name,
                fontSize=14,
                spaceBefore=15,
                spaceAfter=10
            )

            # Normal text style
            normal_style = ParagraphStyle(
                'KoreanNormal',
                parent=styles['Normal'],
                fontName=font_name,
                fontSize=11,
                leading=16,
                alignment=TA_JUSTIFY,
                spaceAfter=8
            )

            # Citation style
            citation_style = ParagraphStyle(
                'Citation',
                parent=styles['Normal'],
                fontName=font_name,
                fontSize=10,
                leading=14,
                leftIndent=10,
                spaceAfter=6
            )

            # Disclaimer style
            disclaimer_style = ParagraphStyle(
                'Disclaimer',
                parent=styles['Normal'],
                fontName=font_name,
                fontSize=9,
                textColor=colors.grey,
                alignment=TA_CENTER,
                spaceBefore=20
            )

            story = []

            # === Document Header ===
            story.append(Paragraph("이혼 소송 준비서면", title_style))
            story.append(Paragraph("(초 안)", ParagraphStyle(
                'Subtitle',
                parent=normal_style,
                alignment=TA_CENTER,
                fontSize=12,
                spaceAfter=30
            )))

            # === Case Information Table ===
            case_data = [
                ["사 건 명", case.title],
                ["생성일시", draft_response.generated_at.strftime('%Y년 %m월 %d일 %H:%M')],
            ]
            case_table = Table(case_data, colWidths=[80, 350])
            case_table.setStyle(TableStyle([
                ('FONTNAME', (0, 0), (-1, -1), font_name),
                ('FONTSIZE', (0, 0), (-1, -1), 11),
                ('ALIGN', (0, 0), (0, -1), 'RIGHT'),
                ('ALIGN', (1, 0), (1, -1), 'LEFT'),
                ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
                ('TOPPADDING', (0, 0), (-1, -1), 8),
            ]))
            story.append(case_table)
            story.append(Spacer(1, 0.4 * inch))

            # === Main Content ===
            story.append(Paragraph("본 문", heading_style))

            # Split content by paragraphs and add to story
            for paragraph_text in draft_response.draft_text.split("\n\n"):
                cleaned_text = paragraph_text.strip()
                if cleaned_text:
                    # Escape XML special characters
                    cleaned_text = (
                        cleaned_text
                        .replace("&", "&amp;")
                        .replace("<", "&lt;")
                        .replace(">", "&gt;")
                    )
                    story.append(Paragraph(cleaned_text, normal_style))

            # === Citations Section ===
            if draft_response.citations:
                story.append(Spacer(1, 0.3 * inch))
                story.append(Paragraph("인용 증거", heading_style))

                for i, citation in enumerate(draft_response.citations, 1):
                    # Citation header
                    labels_str = ", ".join(citation.labels) if citation.labels else "N/A"
                    citation_header = f"<b>[증거 {i}]</b> (ID: {citation.evidence_id})"
                    story.append(Paragraph(citation_header, citation_style))

                    # Citation details
                    story.append(Paragraph(f"분류: {labels_str}", citation_style))

                    # Citation snippet (escape special characters)
                    snippet = (
                        citation.snippet
                        .replace("&", "&amp;")
                        .replace("<", "&lt;")
                        .replace(">", "&gt;")
                    )
                    story.append(Paragraph(f"내용: {snippet}", citation_style))
                    story.append(Spacer(1, 0.1 * inch))

            # === Disclaimer ===
            story.append(Spacer(1, 0.5 * inch))
            story.append(Paragraph(
                "⚠ 본 문서는 AI가 생성한 초안이며, 변호사의 검토 및 수정이 필수입니다.",
                disclaimer_style
            ))

            # Build PDF
            doc.build(story)
            file_buffer.seek(0)

            safe_title = case.title.replace(" ", "_")[:30]
            filename = f"draft_{safe_title}_{draft_response.generated_at.strftime('%Y%m%d')}.pdf"
            content_type = "application/pdf"

            return file_buffer, filename, content_type

        except ImportError:
            raise ValidationError(
                "PDF export is not available. "
                "Please install reportlab: pip install reportlab. "
                "Alternatively, use DOCX format."
            )

    def _register_korean_font(self, pdfmetrics, TTFont) -> bool:
        """
        Register Korean font for PDF generation

        Tries to find and register a Korean font in this order:
        1. Noto Sans KR (bundled or system)
        2. macOS system fonts (AppleGothic, AppleSDGothicNeo)
        3. Linux system fonts (NanumGothic)
        4. Windows system fonts (Malgun Gothic)

        Args:
            pdfmetrics: reportlab pdfmetrics module
            TTFont: reportlab TTFont class

        Returns:
            bool: True if Korean font was registered, False otherwise
        """
        import os

        # Font search paths
        font_candidates = [
            # Bundled font (if exists in project)
            os.path.join(os.path.dirname(__file__), "..", "fonts", "NotoSansKR-Regular.ttf"),
            # macOS system fonts
            "/System/Library/Fonts/AppleSDGothicNeo.ttc",
            "/System/Library/Fonts/Supplemental/AppleGothic.ttf",
            "/Library/Fonts/NotoSansKR-Regular.ttf",
            # Linux system fonts
            "/usr/share/fonts/truetype/nanum/NanumGothic.ttf",
            "/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc",
            # Windows system fonts
            "C:/Windows/Fonts/malgun.ttf",
            "C:/Windows/Fonts/NotoSansKR-Regular.ttf",
        ]

        for font_path in font_candidates:
            if os.path.exists(font_path):
                try:
                    pdfmetrics.registerFont(TTFont("NotoSansKR", font_path))
                    return True
                except Exception:
                    continue

        # No Korean font found - will use default font
        return False
